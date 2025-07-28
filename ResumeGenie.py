import os
import re
import platform
from collections import Counter
import docx
import streamlit as st
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from PIL import Image
import base64
from io import BytesIO

# --- Load and convert image for banner ---
def load_image():
    image_path = r"C:\Users\shaha\AppData\Local\Programs\Python\Python312\WhatsApp Image 2025-04-15 at 13.25.20_a7718cff.jpg"
    with Image.open(image_path) as img:
        buffered = BytesIO()
        img.save(buffered, format="PNG")
        img_base64 = base64.b64encode(buffered.getvalue()).decode()
    return f"data:image/png;base64,{img_base64}"

# --- Function to extract text from a .docx file ---
def extract_text_from_docx(file):
    try:
        doc = docx.Document(file)
        return ' '.join([para.text for para in doc.paragraphs]).lower()
    except Exception as e:
        st.error(f"❌ Error reading the resume: {e}")
        return ""

# --- Function to extract keywords from text ---
def extract_keywords(text):
    words = re.findall(r'\b\w+\b', text.lower())
    stopwords = set([
        "and", "or", "with", "in", "on", "for", "the", "a", "an", "of", "to",
        "from", "at", "by", "is", "are", "this", "that", "as", "be", "have", "has"
    ])
    keywords = [word for word in words if word not in stopwords and len(word) > 2]
    return keywords

# --- ATS Checker Function ---
def ats_checker():
    st.markdown("## 🧠 ATS Checker")
    resume_file = st.file_uploader("📄 Upload your resume (DOCX format)", type="docx")
    if resume_file:
        resume_text = extract_text_from_docx(resume_file)
        if not resume_text:
            st.error("⚠ Failed to extract resume content.")
            return

        job_description = st.text_area("📝 Paste the job description here")
        if job_description:
            jd_keywords = extract_keywords(job_description)
            matched = [kw for kw in jd_keywords if kw in resume_text]
            missing = [kw for kw in jd_keywords if kw not in resume_text]
            score = round((len(matched) / len(jd_keywords)) * 100) if jd_keywords else 0

            st.subheader("🔍 ATS Analysis Result")
            st.success(f"✅ Matched Keywords ({len(matched)}): {matched}")
            st.error(f"❌ Missing Keywords ({len(missing)}): {missing}")
            st.info(f"📊 ATS Compatibility Score: *{score}/100*")

            st.subheader("📋 Resume Formatting Tips")
            common_sections = ["skills", "experience", "education", "projects", "summary"]
            if any(section in resume_text for section in common_sections):
                st.success("✔ Standard sections found (Skills, Education, etc.)")
            else:
                st.warning("⚠ Consider adding common sections like Skills, Education, Projects.")

            if len(resume_text.split()) < 150:
                st.warning("⚠ Resume might be too short. Consider adding more details.")
            else:
                st.success("✔ Resume length looks good.")

            st.balloons()

# --- Resume Ranker Function ---
def resume_ranker():
    st.markdown("## 🎯 Resume Ranker")
    resume_file = st.file_uploader("📄 Upload your resume (DOCX format)", type="docx")
    if resume_file:
        resume_text = extract_text_from_docx(resume_file)
        if not resume_text:
            st.error("⚠ Failed to extract resume content.")
            return

        job_descriptions = []
        jd1 = st.text_area("📝 Job Description 1")
        jd2 = st.text_area("📝 Job Description 2")
        if jd1.strip() and jd2.strip():
            job_descriptions = [jd1.strip(), jd2.strip()]
        else:
            st.warning("⚠ Please provide at least 2 job descriptions.")
            return

        st.subheader("🔍 Ranking Resume against Job Descriptions...")
        vectorizer = TfidfVectorizer()
        rankings = []
        for idx, jd in enumerate(job_descriptions):
            tfidf_matrix = vectorizer.fit_transform([resume_text, jd])
            score = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            rankings.append((idx + 1, score * 100))

        rankings.sort(key=lambda x: x[1], reverse=True)
        for rank, (jd_no, score) in enumerate(rankings, start=1):
            st.write(f"🏆 Rank {rank} → Job Description {jd_no} → Match Score: *{score:.2f}%*")

        best_match = rankings[0]
        st.success(f"✅ Best Match: Job Description {best_match[0]} with a score of {best_match[1]:.2f}%")

# --- Resume Enhancer Function ---
def resume_enhancer():
    st.markdown("## 💡 Resume Enhancer")
    resume_file = st.file_uploader("📄 Upload your resume (DOCX format)", type="docx")
    if resume_file:
        resume_text = extract_text_from_docx(resume_file)
        if not resume_text:
            st.error("⚠ Failed to extract resume content.")
            return

        jd_text = st.text_area("📝 Paste the Job Description here")
        if jd_text:
            jd_keywords = extract_keywords(jd_text)
            resume_keywords = extract_keywords(resume_text)
            jd_keyword_freq = Counter(jd_keywords)
            missing_keywords = [word for word in jd_keyword_freq if word not in resume_keywords]
            top_missing = sorted(missing_keywords, key=lambda x: -jd_keyword_freq[x])[:10]

            st.subheader("🔍 Keyword Match Analysis")
            if top_missing:
                st.info("💡 Based on the job description, here are the TOP 10 keywords to consider adding:")
                for i, kw in enumerate(top_missing, 1):
                    st.write(f"{i}. {kw}")
            else:
                st.success("✅ Great job! Your resume already includes the main keywords.")

# --- Cover Letter Generator ---
def generate_cover_letter():
    st.markdown("## 📝 Cover Letter Generator")
    name = st.text_input("Enter your full name:")
    job_title = st.text_input("Enter the job title:")
    company = st.text_input("Enter the company name:")

    if name and job_title and company:
        cover_letter = f"""
Dear Hiring Manager at {company},

I am writing to express my strong interest in the position of {job_title} at your esteemed organization. With a solid background in relevant technical skills and a strong commitment to continuous learning, I believe I am well-suited for this role.

Upon reviewing the job description, I noticed a strong emphasis on data analysis, problem-solving, machine learning, data science expertise, and job-specific responsibilities. I have hands-on experience and a deep passion for these areas, which closely align with the requirements outlined for this role.

My resume reflects key achievements and contributions in data-driven projects, demonstrating my ability to derive actionable insights and contribute meaningfully to innovative solutions. I am confident that my technical proficiency and enthusiasm for data science make me a valuable candidate for this opportunity.

I would welcome the chance to further discuss how my skills and experience can contribute to {company}’s continued success. Thank you for considering my application. I look forward to the possibility of joining your dynamic team.

Sincerely,
{name}
"""
        st.text_area("Generated Cover Letter", cover_letter.strip(), height=300)

        if st.button("📥 Download Cover Letter"):
            filename = os.path.join(os.path.expanduser('~'), 'Documents', 'Cover_Letter.docx')
            doc = docx.Document()
            doc.add_paragraph(cover_letter.strip())
            try:
                doc.save(filename)
                st.success(f"✅ Cover Letter saved at: {filename}")
            except Exception as e:
                st.error(f"❌ Error saving cover letter: {e}")

# --- Main Menu ---
def main():
    banner = load_image()
    st.markdown(f"""
        <div style="text-align:center;">
            <img src="{banner}" alt="Banner" style="width:100%; max-height:300px; object-fit:cover; border-radius:15px; box-shadow:0 4px 12px rgba(0,0,0,0.3);" />
        </div>
        <h1 style="text-align:center; color:#4A90E2;">ResumeGenie.AI</h1>
        <p style="text-align:center; font-size:18px;">✨ Smart Tools for Resume, ATS, and Cover Letters</p>
        <hr style="border-top: 1px solid #bbb;">
    """, unsafe_allow_html=True)

    menu = ["ATS Checker", "Resume Ranker", "Resume Enhancer", "Generate Cover Letter"]
    choice = st.sidebar.radio("📌 Choose a Tool", menu)

    if choice == "ATS Checker":
        ats_checker()
    elif choice == "Resume Ranker":
        resume_ranker()
    elif choice == "Resume Enhancer":
        resume_enhancer()
    elif choice == "Generate Cover Letter":
        generate_cover_letter()

# --- Run App ---
if __name__ == "__main__":
    main()
